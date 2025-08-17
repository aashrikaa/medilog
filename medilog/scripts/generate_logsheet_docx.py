from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime
import os


def add_heading(document: Document, text: str):
	heading = document.add_heading(text, level=1)
	heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
	return heading


def add_logo(document: Document, logo_path: str):
	if os.path.exists(logo_path):
		p = document.add_paragraph()
		r = p.add_run()
		try:
			r.add_picture(logo_path, width=Inches(1.0))
		except Exception:
			pass


def add_subtitle(document: Document, text: str):
	p = document.add_paragraph()
	r = p.add_run(text)
	r.font.size = Pt(11)
	return p


def set_header_row_style(row):
	for cell in row.cells:
		# Bold text
		for paragraph in cell.paragraphs:
			if paragraph.runs:
				for run in paragraph.runs:
					run.font.bold = True
			else:
				r = paragraph.add_run()
				r.font.bold = True
		# Shading
		tc = cell._tc
		tcPr = tc.get_or_add_tcPr()
		shd = OxmlElement('w:shd')
		shd.set(qn('w:fill'), 'D9EAF7')  # light blue
		tcPr.append(shd)


def set_column_widths(table):
	try:
		table.autofit = False
	except Exception:
		pass
	col_widths = [Inches(1.5), Inches(1.8), Inches(2.7), Inches(0.5)]
	for idx, width in enumerate(col_widths):
		for cell in table.columns[idx].cells:
			cell.width = width


def build_table(document: Document):
	table = document.add_table(rows=1, cols=4)
	table.style = 'Table Grid'
	hdr_cells = table.rows[0].cells
	hdr_cells[0].text = 'Week'
	hdr_cells[1].text = 'Activity Planned'
	hdr_cells[2].text = 'Activity Done'
	hdr_cells[3].text = 'Time Spent (hrs)'
	set_header_row_style(table.rows[0])

	rows = [
		(
			"Week 11 (13–19 July)",
			"Implement dual calendar (AD/BS) and reminder scheduling",
			"Implemented reminders (CRUD, API, UI) with MongoDB; added BS conversion utilities in lib/calendar.ts. "
			"Set up client-side periodic checks for upcoming reminders and Browser Notification API. "
			"node-cron dependency added for future use, but not used in current codebase.",
			"10",
		),
		(
			"Week 12–13 (20 July – 2 Aug)",
			"File upload enhancements & AI integration",
			"Integrated react-dropzone for drag-and-drop (single-file upload). Implemented server-side text extraction for PDF (pdf-parse), "
			"DOCX (mammoth), TXT, and image OCR (tesseract.js). Attempted Python AI integration (blocked by env/deps); "
			"pivoted to Google Gemini and successfully integrated it to extract lab values, generate summaries, and suggest tags for 'Lab Reports' in the upload flow.",
			"20",
		),
		(
			"Week 14 (3–9 Aug)",
			"Document categorization + health insights",
			"Category selection (Lab Reports, Prescriptions, Imaging, Other). Advanced filtering/search (category, tags, date range, lab-only) wired to /api/documents. "
			"Implemented SVG trend charts and abnormal value alerts in components/HealthInsights.tsx.",
			"13",
		),
		(
			"Week 15 (10–16 Aug)",
			"User profile & language support",
			"Added user profile fields in schema (blood type, allergies, emergency contacts); UI integration partial. "
			"QR health summary planned but not fully implemented. English/Nepali labels added to key components and BS date display via lib/calendar.ts — implemented but pending full testing.",
			"12",
		),
		(
			"Week 16 (17 Aug)",
			"Final polish & testing",
			"Responsive UI refinements; improved error handling/logging; partial end-to-end testing; fixed critical bugs; prepared final report. "
			"Note: persistent audit logs are not implemented (console logs only).",
			"14",
		),
	]

	for week, planned, done, hours in rows:
		row_cells = table.add_row().cells
		row_cells[0].text = week
		row_cells[1].text = planned
		row_cells[2].text = done
		row_cells[3].text = hours
		for p in row_cells[3].paragraphs:
			p.alignment = WD_ALIGN_PARAGRAPH.CENTER

	set_column_widths(table)
	return table


def add_footer(document: Document, name: str):
	section = document.sections[-1]
	footer = section.footer
	p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
	p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
	run = p.add_run(f"{name} • Generated on {datetime.now().strftime('%Y-%m-%d %H:%M')}")
	run.font.size = Pt(9)


def main():
	name_for_footer = os.environ.get('MEDILOG_AUTHOR', 'MediLog')
	logo_candidates = [
		'/workspace/medilog/public/logo.png',
		'/workspace/medilog/public/favicon.ico',
		'/workspace/medilog/favicon.ico',
	]

	document = Document()
	for logo in logo_candidates:
		add_logo(document, logo)
	add_heading(document, 'MediLog Logsheet (Continuation)')
	add_subtitle(document, 'Project: MediLog')
	add_subtitle(document, 'Note: Times and statuses reflect corrected entries as discussed.')
	document.add_paragraph()
	build_table(document)
	add_footer(document, name_for_footer)

	output_docx = '/workspace/medilog/MediLog_Logsheet_Continuation.docx'
	document.save(output_docx)
	print(f'DOCX generated at: {output_docx}')

	# Optional PDF export via docx2pdf if available
	try:
		from docx2pdf import convert
		output_pdf = '/workspace/medilog/MediLog_Logsheet_Continuation.pdf'
		convert(output_docx, output_pdf)
		print(f'PDF generated at: {output_pdf}')
	except Exception as e:
		print('PDF export skipped (docx2pdf not available or failed):', e)


if __name__ == '__main__':
	main()